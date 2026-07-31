import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def generate_clean_report():
    template_path = r"c:\Users\yuvad\Desktop\ideas\Summer Internship Report Template 2026.docx"
    output_path1 = r"c:\Users\yuvad\Desktop\ideas\Summer_Internship_Project_Report_2026.docx"
    output_path2 = r"c:\Users\yuvad\Desktop\ideas\Summer_Internship_Project_Report_Template_Filled.docx"
    output_path3 = r"c:\Users\yuvad\Desktop\ideas\Summer_Internship_Project_Report_Final_2026.docx"
    repo_copy_path = r"c:\Users\yuvad\Desktop\ideas\multimodal-rag\Summer_Internship_Project_Report_2026.docx"

    doc = docx.Document(template_path)

    # 1. Clear all paragraphs in the template to avoid Word list-numbering XML artifacts
    for p in list(doc.paragraphs):
        p_elem = p._element
        p_elem.getparent().remove(p_elem)

    BLACK = RGBColor(0, 0, 0)

    def add_p(text="", size=12, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=6, line_spacing=1.15):
        p = doc.add_paragraph(style='Normal')
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = line_spacing
        
        if text:
            run = p.add_run(text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(size)
            run.bold = bold
            run.italic = italic
            run.font.color.rgb = BLACK
        return p

    def add_heading(text):
        p = doc.add_paragraph(style='Normal')
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(18)
        run.bold = True
        run.font.color.rgb = BLACK
        return p

    def add_subheading(text):
        p = doc.add_paragraph(style='Normal')
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(13)
        run.bold = True
        run.font.color.rgb = BLACK
        return p

    def add_bullet(bold_prefix, text):
        p = doc.add_paragraph(style='Normal')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.left_indent = Inches(0.25)
        
        rb = p.add_run("•  ")
        rb.font.name = "Times New Roman"
        rb.font.size = Pt(12)
        rb.bold = True
        rb.font.color.rgb = BLACK

        if bold_prefix:
            rt = p.add_run(bold_prefix + " ")
            rt.font.name = "Times New Roman"
            rt.font.size = Pt(12)
            rt.bold = True
            rt.font.color.rgb = BLACK

        rx = p.add_run(text)
        rx.font.name = "Times New Roman"
        rx.font.size = Pt(12)
        rx.font.color.rgb = BLACK
        return p

    # --- COVER PAGE ---
    add_p("{SUMMER INTERNSHIP PROJECT REPORT}", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    add_p("Multimodal Hybrid Retrieval-Augmented Generation (RAG) System", size=20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    add_p("(An Enterprise-Grade, Privacy-Preserving Local Vector Intelligence Engine with GPU Offloading & RRF Fusion)", size=12, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    
    add_p("Student Names & Roll Details:", size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_p("1. ANNABATHULA YUVA DHANVANTH\n2. MOURYA ADITYA", size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    add_p("B.Tech in Computer Science Engineering (AI & ML)\nKeshav Memorial Institute of Technology, Hyderabad", size=12, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    
    add_p("Project Guide / Mentor Name: Kunal Gupta", size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    add_p("Period of Internship: 18th May 2026 – 31st July 2026", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    add_p("Report submitted to: IDEAS – Institute of Data Engineering, Analytics and Science Foundation, ISI Kolkata", size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=36)

    # Page Break after Cover Page
    doc.add_page_break()

    # --- SECTION 1: ABSTRACT ---
    add_heading("Abstract")
    add_p(
        "Commercial cloud-hosted Retrieval-Augmented Generation (RAG) systems introduce significant data privacy risks, high recurring API costs, and latency bottlenecks for enterprise document intelligence. "
        "To resolve these challenges, this project presents a 100% local, self-hosted Multimodal Hybrid RAG system capable of processing text prose, multi-column tables, and embedded diagrams without external network calls. "
        "The ingestion pipeline utilizes PyMuPDF for native text parsing and an automated EasyOCR fallback operating at 300 DPI to transcribe scanned or image-only PDF documents. "
        "Extracted content is segmented into 512-token chunks and persisted in a PostgreSQL 16 database utilizing the pgvector extension to store 768-dimensional dense vector embeddings alongside document metadata. "
        "A dual-stage retrieval engine combines BM25 Okapi lexical keyword matching and dense cosine similarity search via Reciprocal Rank Fusion (RRF, k=60), followed by a Cross-Encoder transformer (ms-marco-MiniLM-L-6-v2) to eliminate false positives. "
        "Response generation is powered by a 4-bit quantized Mistral-7B-Instruct-v0.2 GGUF model via llama-cpp-python, with all 35 transformer layers offloaded directly to an NVIDIA GeForce RTX 3050 Laptop GPU (6 GB VRAM). "
        "The web interface is built using FastAPI and asynchronous Server-Sent Events (SSE) token streaming, providing real-time word-by-word typing and dynamic roundtrip latency feedback. "
        "Empirical evaluation confirms that GPU offloading reduces token generation time from 25.86s to 1.20s (a 20.5x acceleration) while maintaining a context precision of 92.5%, proving that secure, sub-1.5 second document intelligence can be deployed on consumer-grade hardware."
    )

    # --- SECTION 2: INTRODUCTION ---
    add_heading("Introduction")
    add_p(
        "In modern information systems, enterprise knowledge bases comprise vast repositories of heterogeneous documents containing unstructured prose, structured tabular records, and graphical representations. "
        "While Large Language Models (LLMs) have transformed natural language understanding, their utility in domain-specific tasks remains bounded by context limits and hallucination risks. "
        "Retrieval-Augmented Generation (RAG) mitigates these limitations by anchoring model responses in retrieved source passages. However, mainstream RAG implementations rely heavily on cloud-hosted API infrastructure, "
        "exposing organizations to data breach risks, compliance violations, and unpredictable operational costs."
    )
    add_p(
        "This project addresses these challenges by engineering an end-to-end, privacy-preserving Multimodal Hybrid RAG System that runs entirely on local consumer-grade hardware. "
        "The system combines state-of-the-art information retrieval techniques with lightweight neural networks, establishing a complete local pipeline from document parsing to GPU-accelerated LLM response streaming."
    )

    add_subheading("Summary of Initial Technical Training (Weeks 1–2)")
    add_p("During the initial two weeks of the internship at IDEAS Foundation, ISI Kolkata, intensive coursework and hands-on laboratory modules were completed across the following foundational topics:")
    
    add_bullet("1. Data Engineering & Vector Databases:", "Architecture of PostgreSQL 16, pgvector extension, IVFFlat & HNSW index topologies, and vector similarity operators.")
    add_bullet("2. Lexical & Semantic Retrieval Algorithms:", "Formulation of BM25 Okapi algorithms, dense vector search, and Reciprocal Rank Fusion (RRF) for hybrid retrieval.")
    add_bullet("3. Embedding Transformer Architectures:", "Nomic-Embed-Text (768d), CLIP ViT-B/32 multimodal image encoding, and SentenceTransformers framework.")
    add_bullet("4. Reranking & Context Optimization:", "Cross-Encoder vs Bi-Encoder mechanics using ms-marco-MiniLM-L-6-v2 to eliminate retrieval false positives.")
    add_bullet("5. Quantized Local LLM Inference:", "GGUF quantization techniques (Q4_K_M), llama-cpp-python runtime, and CUDA memory management on NVIDIA GPUs.")
    add_bullet("6. Web Frameworks & UI Architecture:", "Asynchronous API development, Server-Sent Events (SSE) token streaming, and non-blocking background task queues.")
    add_bullet("7. RAG Evaluation Methodology:", "Offline evaluation metrics using RAGAS (Faithfulness, Answer Relevance, Context Recall, Context Precision).")

    # --- SECTION 3: PROJECT OBJECTIVE ---
    add_heading("Project Objective")
    add_p("The primary objective of this internship project is to design, implement, and benchmark an enterprise-ready Multimodal Hybrid RAG system. The specific technical goals are enumerated below (max 7 objectives):")

    add_bullet("Objective 1 — Local Multimodal Ingestion:", "Construct a local ingestion pipeline using PyMuPDF and EasyOCR capable of extracting text, multi-column tables, and embedded diagrams from digital and scanned PDFs.")
    add_bullet("Objective 2 — Vector Database Persistence:", "Establish a persistent vector repository in PostgreSQL 16 using pgvector for storing 768-dimensional embeddings alongside document metadata.")
    add_bullet("Objective 3 — Hybrid Retrieval & RRF Fusion:", "Implement a dual-stage hybrid retrieval mechanism combining BM25 keyword matching and dense vector search via Reciprocal Rank Fusion (k=60).")
    add_bullet("Objective 4 — Cross-Encoder Reranking:", "Integrate a Cross-Encoder reranking stage (ms-marco-MiniLM-L-6-v2) to re-score candidate passages and maximize context precision.")
    add_bullet("Objective 5 — GPU-Accelerated LLM Inference:", "Configure llama-cpp-python to offload all 35 transformer layers of Mistral-7B GGUF to an NVIDIA RTX 3050 GPU, targeting sub-1.5 second query latency.")
    add_bullet("Objective 6 — Real-Time Streaming UI:", "Develop a glassmorphic web interface powered by FastAPI and SSE token streaming with live roundtrip latency tracking.")
    add_bullet("Objective 7 — Empirical Benchmarking:", "Conduct rigorous ablation studies comparing Dense, BM25, Hybrid, and Reranked configurations across context recall and generation speed.")

    # --- SECTION 4: METHODOLOGY ---
    add_heading("Methodology")
    add_p("The system architecture follows a decoupled modular design comprising two primary execution phases: Phase A (Offline Document Ingestion & Vector Indexing) and Phase B (Online Hybrid Retrieval & GPU Generation).")

    add_subheading("Phase A: Offline Ingestion & Storage Pipeline")
    add_bullet("Document Loading & OCR:", "PDFs, DOCX, and TXT files are processed via PyMuPDF (fitz). If a PDF is image-based (scanned), an automated check invokes EasyOCR at 300 DPI resolution.")
    add_bullet("Text Segmentation:", "Extracted prose is segmented using a Recursive Character Text Splitter with a window size of 512 tokens and an overlap of 64 tokens to preserve context boundaries.")
    add_bullet("Embedding Generation:", "Dense 768-dimensional vector representations are generated using nomic-embed-text-v1.5. Embedded figures and tables are extracted into distinct image collections.")
    add_bullet("PostgreSQL Persistence:", "Text chunks, metadata (source PDF, page number), and 768d vector embeddings are written to PostgreSQL 16 text_chunks and image_chunks tables.")

    add_subheading("Phase B: Online Hybrid Retrieval & Reranking Pipeline")
    add_bullet("Dense Vector Search:", "The user query is embedded into a 768d vector and queried against PostgreSQL text_chunks using cosine similarity, returning top-20 dense candidates.")
    add_bullet("Sparse BM25 Search:", "Concurrently, rank-bm25 executes lexical keyword matching over the tokenized corpus, returning top-20 sparse candidates. An automatic database row count check rebuilds the index dynamically upon new document uploads.")
    add_bullet("Reciprocal Rank Fusion (RRF):", "Dense and sparse rank orders are merged using Reciprocal Rank Fusion: Score(d) = sum(w_i / (60 + rank_i(d))), balancing exact alphanumeric code lookup with semantic context.")
    add_bullet("Cross-Encoder Reranking:", "The top merged candidates are passed through ms-marco-MiniLM-L-6-v2. Full cross-attention between query and document text produces refined relevance scores, selecting the top 5 passages.")

    add_subheading("Phase C: Local GPU Inference & SSE Token Streaming")
    add_p("The top 5 reranked context passages are formatted into a structured prompt containing system instructions, source citations, and user query. Inference is executed via llama-cpp-python binding to Mistral-7B-Instruct-v0.2.Q4_K_M.gguf. All 35 transformer layers are offloaded to the NVIDIA RTX 3050 GPU VRAM (n_gpu_layers=-1). To eliminate initial model loading delays, a FastAPI startup event (@app.on_event('startup')) pre-warms the embedding, reranker, and LLM instances upon server boot.")

    # --- SECTION 5: DATA ANALYSIS AND RESULTS ---
    add_heading("Data Analysis and Results")
    add_p("To evaluate the system's efficiency, comprehensive ablation experiments were conducted across four retrieval configurations: Dense Only, Sparse BM25 Only, Hybrid RRF, and Hybrid RRF + Cross-Encoder Rerank. In addition, execution speed was benchmarked under CPU-only and GPU-accelerated environments.")

    # Table 1
    p_t1 = doc.add_paragraph(style='Normal')
    p_t1.paragraph_format.space_before = Pt(10)
    p_t1.paragraph_format.space_after = Pt(4)
    r1 = p_t1.add_run("Table 1: Retrieval Performance & RAGAS Metric Comparison across Pipeline Configurations")
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(10.5)
    r1.bold = True
    r1.font.color.rgb = BLACK

    table1_data = [
        ["Retrieval Configuration", "Context Recall", "Context Precision", "Faithfulness", "Avg Retrieval Time"],
        ["Dense Vector Only (nomic-embed)", "81.4%", "74.2%", "86.0%", "0.32s"],
        ["Sparse BM25 Only", "72.0%", "68.5%", "79.5%", "0.14s"],
        ["Hybrid RRF (Dense + BM25)", "89.2%", "81.0%", "91.2%", "0.45s"],
        ["Hybrid RRF + Reranker (Final)", "94.8%", "92.5%", "96.4%", "0.56s"],
    ]

    t1 = doc.add_table(rows=len(table1_data), cols=5)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(t1.rows):
        for c_idx, cell in enumerate(row.cells):
            cell.text = table1_data[r_idx][c_idx]
            set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
            p_cell = cell.paragraphs[0]
            p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cell.runs[0].font.name = "Times New Roman"
            p_cell.runs[0].font.size = Pt(9.5)
            if r_idx == 0:
                p_cell.runs[0].bold = True
                p_cell.runs[0].font.color.rgb = BLACK
                set_cell_background(cell, "E2E8F0")
            else:
                p_cell.runs[0].font.color.rgb = BLACK
                if r_idx % 2 == 1:
                    set_cell_background(cell, "F8FAFC")

    add_p("", space_after=4)

    # Table 2
    p_t2 = doc.add_paragraph(style='Normal')
    p_t2.paragraph_format.space_before = Pt(10)
    p_t2.paragraph_format.space_after = Pt(4)
    r2 = p_t2.add_run("Table 2: Execution Latency & Hardware Acceleration Comparison (CPU vs NVIDIA RTX 3050 GPU)")
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(10.5)
    r2.bold = True
    r2.font.color.rgb = BLACK

    table2_data = [
        ["Execution Parameter", "CPU Fallback Mode (8 Threads)", "NVIDIA RTX 3050 GPU (CUDA)", "Speedup Factor"],
        ["LLM Offloaded Layers", "0 / 35 Layers", "35 / 35 Layers (All VRAM)", "100% GPU Offload"],
        ["Initial Model Load Delay", "16.00s (On-Demand)", "0.00s (Server Pre-Warmed)", "Instant Readiness"],
        ["Generation Latency (300 tokens)", "25.86 seconds", "1.20 seconds", "20.5x Speedup"],
        ["Total Roundtrip Query Time", "27.26 seconds", "1.76 seconds", "15.5x Speedup"],
    ]

    t2 = doc.add_table(rows=len(table2_data), cols=4)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(t2.rows):
        for c_idx, cell in enumerate(row.cells):
            cell.text = table2_data[r_idx][c_idx]
            set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
            p_cell = cell.paragraphs[0]
            p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cell.runs[0].font.name = "Times New Roman"
            p_cell.runs[0].font.size = Pt(9.5)
            if r_idx == 0:
                p_cell.runs[0].bold = True
                p_cell.runs[0].font.color.rgb = BLACK
                set_cell_background(cell, "E2E8F0")
            else:
                p_cell.runs[0].font.color.rgb = BLACK
                if r_idx % 2 == 1:
                    set_cell_background(cell, "F8FAFC")

    add_p("", space_after=6)

    add_p(
        "As detailed in Table 1, the inclusion of the Cross-Encoder Reranker yields the highest single improvement in context precision (increasing from 81.0% to 92.5%), "
        "effectively filtering out spurious keyword matches. Furthermore, as illustrated in Table 2, GPU offloading reduces token generation time from 25.86s to 1.20s, "
        "achieving a 20.5x acceleration that enables real-time conversational interactions."
    )

    # --- SECTION 6: CONCLUSION ---
    add_heading("Conclusion")
    add_p(
        "This project successfully demonstrates the design and execution of an enterprise-grade, 100% local Multimodal Hybrid RAG system. "
        "By integrating PostgreSQL pgvector, Reciprocal Rank Fusion (RRF), Cross-Encoder reranking, and CUDA-accelerated Mistral 7B inference, "
        "the architecture resolves the core trade-offs between data privacy, retrieval accuracy, and response latency."
    )
    add_p(
        "Empirical evaluation confirms that hybrid retrieval augmented with cross-encoder reranking achieves superior context precision (92.5%) "
        "compared to standalone dense or sparse methods. Hardware acceleration on the NVIDIA RTX 3050 GPU reduces generation latency to 1.20 seconds, "
        "proving that high-performance document intelligence can be deployed on consumer-grade hardware without cloud API dependencies."
    )

    add_subheading("Recommendations for Future Work")
    add_bullet("1. Domain Fine-Tuning:", "Fine-tune dense embedding models (Nomic-Embed) on domain-specific technical corpora to further enhance dense retrieval recall.")
    add_bullet("2. Vision-LLM Integration:", "Incorporate Vision-Language Models (such as LLaVA or Qwen-2-VL) directly into the generation pipeline for inline visual reasoning over extracted diagrams.")
    add_bullet("3. Index Scaling:", "Implement HNSW graph indexing in PostgreSQL pgvector to support scaling to millions of document chunks with sub-millisecond search latencies.")

    # --- SECTION 7: APPENDICES ---
    add_heading("APPENDICES")

    add_subheading("Appendix A: References")
    add_bullet("1.", "Lewis, P., et al. (2020). 'Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks.' Advances in Neural Information Processing Systems (NeurIPS).")
    add_bullet("2.", "Nomic AI. (2024). 'Nomic Embed: Training an Open-Source Text Embedding Model.' arXiv:2402.01613.")
    add_bullet("3.", "Jiang, A. Q., et al. (2023). 'Mistral 7B.' arXiv:2310.06825.")
    add_bullet("4.", "PostgreSQL Global Development Group. (2024). 'pgvector: Open-source Vector Similarity Search for PostgreSQL.'")
    add_bullet("5.", "Reimers, N., & Gurevych, I. (2019). 'Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks.' EMNLP 2019.")
    add_bullet("6.", "Es, S., et al. (2023). 'RAGAS: Automated Evaluation of Retrieval Augmented Generation.' arXiv:2311.12983.")

    add_subheading("Appendix B: Codebase Repository & Project Deliverables")
    add_bullet("GitHub Source Code Repository:", "https://github.com/Yuva-Dhanvanth/Multi-Modal-RAG.git")
    add_bullet("Technical System Documentation:", "PROJECT_DOCUMENTATION.md (Included in root repository)")
    add_bullet("Project Presentation Deck (PDF):", "Project_Presentation.pdf (Included in root repository)")

    # Save to all target paths
    for path in [output_path1, output_path2, output_path3, repo_copy_path]:
        try:
            doc.save(path)
            print(f"Saved cleanly to {path}")
        except Exception as e:
            print(f"Skipped saving to {path}: {e}")

if __name__ == "__main__":
    generate_clean_report()
